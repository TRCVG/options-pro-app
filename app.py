import streamlit as st
from jugaad_data.nse import derivatives_df, index_df
import datetime
import calendar
import pandas as pd
import plotly.graph_objects as go
import time
import io
import copy
from docx import Document
from docx.shared import Inches

# ==========================================
# PAGE CONFIGURATION
# ==========================================
st.set_page_config(layout="wide", page_title="NSE Options Pro Terminal")

# --- INVISIBLE TOP ANCHOR FOR SCROLL BUTTON ---
st.markdown("<div id='top'></div>", unsafe_allow_html=True)

# ==========================================
# TOP SWITCH DISPLAY (TOGGLE BETWEEN MODES)
# ==========================================
st.markdown("### 🎛️ Terminal Workspace Switcher")
is_batch_mode = st.toggle("⚡ Enable Multi-Task Batch Terminal (No Charts)", value=True)

st.markdown("---")

if is_batch_mode:
    st.title("⚡ Multi-Task Batch Strategy Terminal")
    st.markdown("Queue up multiple independent tasks below. You can clone, reorder, minimize, or maximize tasks to manage your workspace efficiently.")
    
    # Initialize session state for batch tasks
    if 'batch_tasks' not in st.session_state:
        st.session_state.batch_tasks = [{
            "id": 0,
            "minimized": False,
            "series": [{
                "series_id": 0,
                "symbol": "NIFTY",
                "option_type": "CE",
                "strikes": "24000",
                "expiry": datetime.date.today(),
                "from_dt": datetime.date.today() - datetime.timedelta(days=10),
                "to_dt": datetime.date.today(),
                "lot_size": 1,
                "action": "Buy (Long)"
            }]
        }]
        st.session_state.batch_next_task_id = 1

    # --- GLOBAL MINIMIZE / MAXIMIZE ALL BUTTONS ---
    ctrl_col1, ctrl_col2, ctrl_col3 = st.columns([2, 2, 6])
    with ctrl_col1:
        if st.button("🗗 Minimize All Tasks"):
            for t in st.session_state.batch_tasks:
                t["minimized"] = True
            st.rerun()
    with ctrl_col2:
        if st.button("🗖 Maximize All Tasks"):
            for t in st.session_state.batch_tasks:
                t["minimized"] = False
            st.rerun()

    st.markdown("")

    # --- RENDER DYNAMIC TASK CARDS ---
    for task_idx, task in enumerate(st.session_state.batch_tasks):
        
        # Build concise summary label string (e.g. NIFTY 24000 CE (B) 21/06)
        summary_parts = []
        for s in task.get('series', []):
            action_code = "(B)" if "Buy" in s["action"] else "(S)"
            exp_str = s["expiry"].strftime("%d/%m") if isinstance(s["expiry"], datetime.date) else str(s["expiry"])
            summary_parts.append(f"{s['symbol']} {s['strikes']} {s['option_type']} {action_code} {exp_str}")
        summary_label = " + ".join(summary_parts) if summary_parts else "Empty Task"

        # Check if task is minimized
        if task.get("minimized", False):
            # Render Minimized Square / Card View
            with st.container(border=True):
                m_col1, m_col2, m_col3 = st.columns([7, 1.5, 1.5])
                with m_col1:
                    st.markdown(f"**📌 Task #{task_idx + 1} | {summary_label}**")
                with m_col2:
                    if st.button("🗖 Maximize", key=f"max_task_{task_idx}_{task['id']}"):
                        task["minimized"] = False
                        st.rerun()
                with m_col3:
                    if len(st.session_state.batch_tasks) > 1:
                        if st.button("🗑️ Delete", key=f"del_min_task_{task_idx}_{task['id']}"):
                            st.session_state.batch_tasks.pop(task_idx)
                            st.rerun()
        else:
            # Render Expanded Task Card
            with st.container(border=True):
                t_col_h1, t_col_h2, t_col_h3, t_col_h4, t_col_h5 = st.columns([3.5, 1.5, 1.8, 1.6, 1.6])
                with t_col_h1:
                    st.subheader(f"📌 Task #{task_idx + 1}")
                with t_col_h2:
                    if st.button("📋 Clone", key=f"clone_task_{task_idx}_{task['id']}"):
                        duplicated_task = copy.deepcopy(task)
                        duplicated_task['id'] = st.session_state.batch_next_task_id
                        st.session_state.batch_next_task_id += 1
                        st.session_state.batch_tasks.insert(task_idx + 1, duplicated_task)
                        st.rerun()
                with t_col_h3:
                    col_mu, col_md = st.columns(2)
                    with col_mu:
                        if task_idx > 0 and st.button("⬆️", key=f"move_up_{task_idx}_{task['id']}"):
                            st.session_state.batch_tasks[task_idx], st.session_state.batch_tasks[task_idx - 1] = st.session_state.batch_tasks[task_idx - 1], st.session_state.batch_tasks[task_idx]
                            st.rerun()
                    with col_md:
                        if task_idx < len(st.session_state.batch_tasks) - 1 and st.button("⬇️", key=f"move_down_{task_idx}_{task['id']}"):
                            st.session_state.batch_tasks[task_idx], st.session_state.batch_tasks[task_idx + 1] = st.session_state.batch_tasks[task_idx + 1], st.session_state.batch_tasks[task_idx]
                            st.rerun()
                with t_col_h4:
                    if st.button("🗗 Minimize", key=f"min_task_{task_idx}_{task['id']}"):
                        task["minimized"] = True
                        st.rerun()
                with t_col_h5:
                    if len(st.session_state.batch_tasks) > 1:
                        if st.button("🗑️ Delete", key=f"del_task_{task_idx}_{task['id']}"):
                            st.session_state.batch_tasks.pop(task_idx)
                            st.rerun()

                # Render multiple series legs inside this task
                if 'series' not in task:
                    task['series'] = [{
                        "series_id": 0,
                        "symbol": "NIFTY",
                        "option_type": "CE",
                        "strikes": "24000",
                        "expiry": datetime.date.today(),
                        "from_dt": datetime.date.today() - datetime.timedelta(days=10),
                        "to_dt": datetime.date.today(),
                        "lot_size": 1,
                        "action": "Buy (Long)"
                    }]

                for s_idx, series_item in enumerate(task['series']):
                    st.markdown(f"**--- Series Leg #{s_idx + 1} ---**")
                    
                    sc1, sc2, sc3 = st.columns([2.5, 8, 1.5])
                    with sc1:
                        if st.button("➕ Add Another Series", key=f"t_{task_idx}_add_s_{series_item['series_id']}"):
                            new_s_id = max([s['series_id'] for s in task['series']], default=0) + 1
                            task['series'].append({
                                "series_id": new_s_id,
                                "symbol": series_item["symbol"],
                                "option_type": series_item["option_type"],
                                "strikes": series_item["strikes"],
                                "expiry": series_item["expiry"],
                                "from_dt": series_item["from_dt"],
                                "to_dt": series_item["to_dt"],
                                "lot_size": 1,
                                "action": series_item["action"]
                            })
                            st.rerun()
                    with sc2:
                        pass
                    with sc3:
                        if len(task['series']) > 1:
                            if st.button("❌ Remove Series", key=f"t_{task_idx}_rem_s_{series_item['series_id']}"):
                                task['series'].pop(s_idx)
                                st.rerun()

                    c1, c2, c3 = st.columns(3)
                    with c1:
                        series_item["symbol"] = st.selectbox("Symbol", ["NIFTY", "BANKNIFTY", "FINNIFTY"], index=["NIFTY", "BANKNIFTY", "FINNIFTY"].index(series_item["symbol"]), key=f"b_sym_{task_idx}_{series_item['series_id']}")
                        series_item["option_type"] = st.selectbox("Option Type", ["CE", "PE"], index=["CE", "PE"].index(series_item["option_type"]), key=f"b_opt_{task_idx}_{series_item['series_id']}")
                    with c2:
                        series_item["strikes"] = st.text_input("Strike Prices (comma separated)", series_item["strikes"], key=f"b_strk_{task_idx}_{series_item['series_id']}")
                        series_item["lot_size"] = st.number_input("Lot Size", value=series_item["lot_size"], key=f"b_lot_{task_idx}_{series_item['series_id']}")
                    with c3:
                        series_item["action"] = st.selectbox("Action", ["Buy (Long)", "Sell (Short)"], index=0 if "Buy" in series_item["action"] else 1, key=f"b_act_{task_idx}_{series_item['series_id']}")
                        series_item["expiry"] = st.date_input("Expiry Date", value=series_item["expiry"], key=f"b_exp_{task_idx}_{series_item['series_id']}")

                    c4, c5 = st.columns(2)
                    with c4:
                        series_item["from_dt"] = st.date_input("Entry Date (From Date)", value=series_item["from_dt"], key=f"b_from_{task_idx}_{series_item['series_id']}")
                    with c5:
                        series_item["to_dt"] = st.date_input("Exit Date (To Date)", value=series_item["to_dt"], key=f"b_to_{task_idx}_{series_item['series_id']}")
                    
                    st.markdown("")

    # --- ADD NEW TASK BUTTON ---
    if st.button("➕ Add New Task"):
        new_task_id = st.session_state.batch_next_task_id
        st.session_state.batch_next_task_id += 1
        st.session_state.batch_tasks.append({
            "id": new_task_id,
            "minimized": False,
            "series": [{
                "series_id": 0,
                "symbol": "NIFTY",
                "option_type": "CE",
                "strikes": "24000",
                "expiry": datetime.date.today(),
                "from_dt": datetime.date.today() - datetime.timedelta(days=10),
                "to_dt": datetime.date.today(),
                "lot_size": 1,
                "action": "Buy (Long)"
            }]
        })
        st.rerun()

    st.markdown("---")
    
    # Run Batch Execution Button with Comparative Table Summary Output
    if st.button("🚀 Run All Batch Tasks Sequentially", type="primary"):
        st.markdown("### 📊 Batch Execution Summary Table")
        
        table_rows = []
        
        with st.spinner("Running sequential batch calculations across all tasks..."):
            for task_idx, task in enumerate(st.session_state.batch_tasks):
                
                summary_parts = []
                for s in task.get('series', []):
                    action_code = "(B)" if "Buy" in s["action"] else "(S)"
                    exp_str = s["expiry"].strftime("%d/%m") if isinstance(s["expiry"], datetime.date) else str(s["expiry"])
                    summary_parts.append(f"{s['symbol']} {s['strikes']} {s['option_type']} {action_code} {exp_str}")
                task_description = " + ".join(summary_parts) if summary_parts else "Empty Task"

                task_combined_dfs = []
                success_flag = "Success ✅"
                
                for s_item in task['series']:
                    strike_list = [int(s.strip()) for s in s_item["strikes"].split(",") if s.strip().isdigit()]
                    if not strike_list:
                        continue
                        
                    for strike in strike_list:
                        try:
                            df = derivatives_df(
                                symbol=s_item["symbol"],
                                from_date=s_item["from_dt"],
                                to_date=s_item["to_dt"],
                                expiry_date=s_item["expiry"],
                                instrument_type="OPTIDX",
                                strike_price=strike,
                                option_type=s_item["option_type"]
                            )
                            
                            if not df.empty:
                                df['DATE'] = pd.to_datetime(df['DATE'])
                                df = df.sort_values('DATE').reset_index(drop=True)
                                
                                entry_price = (0.125 * df['OPEN'].iloc[0] + 0.125 * df['CLOSE'].iloc[0] + 0.375 * df['HIGH'].iloc[0] + 0.375 * df['LOW'].iloc[0])
                                
                                if s_item["action"] == "Buy (Long)":
                                    df['P&L (Points)'] = df['CLOSE'] - entry_price
                                    pnl_open = (df['OPEN'] - entry_price) * s_item["lot_size"]
                                    pnl_high = (df['HIGH'] - entry_price) * s_item["lot_size"]
                                    pnl_low = (df['LOW'] - entry_price) * s_item["lot_size"]
                                    pnl_close = (df['CLOSE'] - entry_price) * s_item["lot_size"]
                                else:
                                    df['P&L (Points)'] = entry_price - df['CLOSE']
                                    pnl_open = (entry_price - df['OPEN']) * s_item["lot_size"]
                                    pnl_high = (entry_price - df['LOW']) * s_item["lot_size"]
                                    pnl_low = (entry_price - df['HIGH']) * s_item["lot_size"]
                                    pnl_close = (entry_price - df['CLOSE']) * s_item["lot_size"]
                                    
                                df['TOTAL P&L (₹)'] = df['P&L (Points)'] * s_item["lot_size"]
                                df['PNL_OPEN'] = pnl_open
                                df['PNL_HIGH'] = pnl_high
                                df['PNL_LOW'] = pnl_low
                                df['PNL_CLOSE'] = pnl_close
                                
                                task_combined_dfs.append(df[['DATE', 'PNL_OPEN', 'PNL_HIGH', 'PNL_LOW', 'PNL_CLOSE', 'TOTAL P&L (₹)']])
                        except Exception as ex:
                            success_flag = "Error / No Data ⚠️"
                
                if task_combined_dfs:
                    task_master = pd.concat(task_combined_dfs, ignore_index=True)
                    portfolio_pnl = task_master.groupby('DATE').agg({
                        'PNL_OPEN': 'sum',
                        'PNL_HIGH': 'sum',
                        'PNL_LOW': 'sum',
                        'PNL_CLOSE': 'sum',
                        'TOTAL P&L (₹)': 'sum'
                    }).reset_index()

                    portfolio_pnl['ROW_MAX'] = portfolio_pnl[['PNL_OPEN', 'PNL_HIGH', 'PNL_LOW', 'PNL_CLOSE']].max(axis=1)
                    portfolio_pnl['ROW_MIN'] = portfolio_pnl[['PNL_OPEN', 'PNL_HIGH', 'PNL_LOW', 'PNL_CLOSE']].min(axis=1)

                    t_overall_max = portfolio_pnl['ROW_MAX'].max()
                    t_overall_min = portfolio_pnl['ROW_MIN'].min()
                    t_final_closing = portfolio_pnl['TOTAL P&L (₹)'].iloc[-1] if not portfolio_pnl.empty else 0.0

                    table_rows.append({
                        "Task #": f"Task #{task_idx + 1}",
                        "Strategy Description": task_description,
                        "Max High (₹)": round(t_overall_max, 2),
                        "Min Low (₹)": round(t_overall_min, 2),
                        "Final Closing P&L (₹)": round(t_final_closing, 2),
                        "Status": success_flag
                    })
                else:
                    table_rows.append({
                        "Task #": f"Task #{task_idx + 1}",
                        "Strategy Description": task_description,
                        "Max High (₹)": 0.0,
                        "Min Low (₹)": 0.0,
                        "Final Closing P&L (₹)": 0.0,
                        "Status": "No Data Found ❌"
                    })
        
        # Display Comparative Summary Table
        if table_rows:
            summary_df = pd.DataFrame(table_rows)
            st.dataframe(summary_df, use_container_width=True, hide_index=True)

else:
    # ==========================================
    # DYNAMIC DATE & STRIKE CALCULATIONS (ORIGINAL APP)
    # ==========================================
    today = datetime.date.today()

    default_entry = today.replace(day=1)
    default_exit = today

    year = today.year
    month = today.month
    last_day_of_month = calendar.monthrange(year, month)[1]
    last_date = datetime.date(year, month, last_day_of_month)
    offset = (last_date.weekday() - calendar.TUESDAY) % 7
    default_expiry = last_date - datetime.timedelta(days=offset)

    def get_nearby_strikes(symbol, target_date, num_strikes=2):
        idx_map = {
            "NIFTY": ("NIFTY 50", 50),
            "BANKNIFTY": ("NIFTY BANK", 100),
            "FINNIFTY": ("NIFTY FIN SERVICE", 50)
        }
        
        idx_sym, step = idx_map.get(symbol, ("NIFTY 50", 50))
        start_date = target_date - datetime.timedelta(days=7)
        
        try:
            df = index_df(symbol=idx_sym, from_date=start_date, to_date=target_date)
            if not df.empty:
                close_col = next((col for col in df.columns if 'CLOSE' in col.upper()), None)
                date_col = next((col for col in df.columns if 'DATE' in col.upper()), None)
                
                if close_col and date_col:
                    df[date_col] = pd.to_datetime(df[date_col])
                    df = df.sort_values(date_col)
                    latest_close = float(df[close_col].iloc[-1])
                    
                    atm = round(latest_close / step) * step
                    strikes = [atm + (i * step) for i in range(-num_strikes, num_strikes + 1)]
                    return strikes, latest_close
        except Exception as e:
            pass
            
        return None, None

    # ==========================================
    # HEADER
    # ==========================================
    st.title("⚙️ Unified Options Backtester & Price Comparator 1 ")
    st.write("Configure your strikes and trade setup below. A single click will generate both the Price Action chart and the P&L Backtest chart.")

    if 'blocks' not in st.session_state:
        st.session_state.blocks = [0]
        st.session_state.next_id = 1

    st.markdown("---")

    # ==========================================
    # 1. DYNAMIC CONFIGURATION BLOCKS
    # ==========================================
    if len(st.session_state.blocks) == 0:
        if st.button("➕ Add Initial Series"):
            st.session_state.blocks.append(st.session_state.next_id)
            st.session_state.next_id += 1
            st.rerun()

    all_series_inputs = []

    for block_idx, block_id in enumerate(st.session_state.blocks):
        
        h_col1, h_col2, h_col3 = st.columns([2.5, 8, 1.5])
        
        with h_col1:
            if st.button("➕ Add Another Series", key=f"add_btn_{block_id}"):
                new_id = st.session_state.next_id
                keys_to_copy = ["sym", "opt", "smode", "strk", "exp", "from", "to", "lot", "act", "ref"]
                for key in keys_to_copy:
                    old_key = f"{key}_{block_id}"
                    new_key = f"{key}_{new_id}"
                    if old_key in st.session_state:
                        st.session_state[new_key] = st.session_state[old_key]
                        
                st.session_state.blocks.append(new_id)
                st.session_state.next_id += 1
                st.rerun()
                
        with h_col2:
            st.markdown(f"#### 🔍 Option Configuration Block #{block_idx+1}")
            
        with h_col3:
            if len(st.session_state.blocks) > 1:
                if st.button("❌ Remove", key=f"del_btn_{block_id}"):
                    st.session_state.blocks.remove(block_id)
                    st.rerun()
                
        c1, c2, c3 = st.columns(3)
        
        with c1:
            symbol = st.selectbox("Symbol", ["NIFTY", "BANKNIFTY", "FINNIFTY"], key=f"sym_{block_id}")
            option_type = st.selectbox("Option Type", ["CE", "PE"], key=f"opt_{block_id}")
            
            strike_mode = st.radio("Strike Selection Method", ["Manual Entry"], horizontal=True, key=f"smode_{block_id}")
            
            if strike_mode == "Manual Entry":
                default_strikes = "24000" if block_idx == 0 else "24200, 24300"
                strikes_input = st.text_input("Strike Prices (comma separated)", default_strikes, key=f"strk_{block_id}")
                strike_list = [int(s.strip()) for s in strikes_input.split(",") if s.strip().isdigit()]
            else:
                st.info("🤖 Strikes will be automatically calculated based on the Index Close Price on your Entry Date.")
                strike_list = []
            
        with c2:
            expiry_dt = st.date_input("Expiry Date", value=default_expiry, key=f"exp_{block_id}")
            from_dt = st.date_input("Entry Date (From Date)", value=default_entry, key=f"from_{block_id}")
            to_dt = st.date_input("Exit Date (To Date)", value=default_exit, key=f"to_{block_id}")
            
        with c3:
            lot_size = st.number_input(f"Lot Size", min_value=1, value=1, step=1, key=f"lot_{block_id}")
            trade_direction = st.selectbox("Action", ["Buy (Long)", "Sell (Short)"], key=f"act_{block_id}")
            entry_price_ref = st.selectbox("Entry Trigger (Line Chart Only)", ["OPEN", "HIGH", "LOW", "CLOSE"], index=0, key=f"ref_{block_id}")
        
        all_series_inputs.append({
            "symbol": symbol,
            "option_type": option_type,
            "strike_mode": strike_mode,
            "strikes": strike_list,
            "expiry_dt": expiry_dt,
            "from_dt": from_dt,
            "to_dt": to_dt,
            "lot_size": lot_size,
            "trade_direction": trade_direction,
            "entry_price_ref": entry_price_ref
        })
        st.markdown("---")

    # ==========================================
    # 2. GLOBAL SETTINGS (DISPLAY ONLY)
    # ==========================================
    st.markdown("### ⚙️ Global Price Display Preferences")
    st.write("Select how you want to visualize the upper Premium Chart:")

    chart_type = st.radio("Chart Type", ["Line Chart (Custom Price Points)", "Candlestick Chart"], horizontal=True)

    disable_checkboxes = (chart_type == "Candlestick Chart")

    col_o, col_h, col_l, col_c = st.columns(4)
    with col_o: show_open = st.checkbox("Open (O)", value=False, disabled=disable_checkboxes)
    with col_h: show_high = st.checkbox("High (H)", value=False, disabled=disable_checkboxes)
    with col_l: show_low = st.checkbox("Low (L)", value=False, disabled=disable_checkboxes)
    with col_c: show_close = st.checkbox("Close (C)", value=True, disabled=disable_checkboxes) 

    st.markdown("---")

    # ==========================================
    # 3. MASTER FETCH & PROCESS BUTTON
    # ==========================================
    if st.button("🚀 Fetch, Compare & Backtest All", type="primary"):
        
        if len(all_series_inputs) == 0:
            st.error("You need at least one Option Configuration Block to run a backtest!")
            st.stop()
            
        if chart_type == "Line Chart (Custom Price Points)" and not (show_open or show_high or show_low or show_close):
            st.error("Please select at least one price checkbox (Open, High, Low, or Close) to draw the price chart!")
            st.stop()
            
        fig_price = go.Figure()
        fig_growth_bar = go.Figure()
        fig_pnl_individual = go.Figure()
        
        data_found = False
        all_combined_dfs = []
        max_retries = 3
        
        bar_labels = []
        bar_values = []
        bar_colors = []
        
        with st.spinner("Fetching market data, building charts, and calculating complex P&L..."):
            
            def calc_growth(start_val, end_val):
                if start_val == 0: return 0.0, ""
                pct = ((end_val - start_val) / start_val) * 100
                color = "green" if pct >= 0 else "red"
                sign = "+" if pct >= 0 else ""
                return pct, f" <span style='color:{color}'>[{sign}{pct:.1f}%]</span>"

            for block_idx, config in enumerate(all_series_inputs):
                
                for strike in config["strikes"]:
                    label_base = f"Block #{block_idx+1} | {config['symbol']} {strike} {config['option_type']}"
                    action_short = "Buy" if config["trade_direction"] == "Buy (Long)" else "Sell"
                    
                    for attempt in range(max_retries):
                        try:
                            df = derivatives_df(
                                symbol=config["symbol"],
                                from_date=config["from_dt"],
                                to_date=config["to_dt"],
                                expiry_date=config["expiry_dt"],
                                instrument_type="OPTIDX",
                                strike_price=strike,
                                option_type=config["option_type"]
                            )
                            
                            if not df.empty:
                                df['DATE'] = pd.to_datetime(df['DATE'])
                                df = df.sort_values('DATE').reset_index(drop=True)
                                
                                # --- 1. PREMIUM PRICE VISUALIZATION ---
                                if chart_type == "Candlestick Chart":
                                    pct, pct_str = calc_growth(df['CLOSE'].iloc[0], df['CLOSE'].iloc[-1])
                                    name = f"{label_base} (Candle)"
                                    fig_price.add_trace(go.Candlestick(
                                        x=df['DATE'],
                                        open=df['OPEN'],
                                        high=df['HIGH'],
                                        low=df['LOW'],
                                        close=df['CLOSE'],
                                        name=name
                                    ))
                                    bar_labels.append(name)
                                    bar_values.append(pct)
                                    bar_colors.append('rgba(0, 200, 0, 0.7)' if pct >= 0 else 'rgba(255, 0, 0, 0.7)')
                                
                                else:
                                    if show_open: 
                                        pct, pct_str = calc_growth(df['OPEN'].iloc[0], df['OPEN'].iloc[-1])
                                        name = f"{label_base} (O)"
                                        fig_price.add_trace(go.Scatter(x=df['DATE'], y=df['OPEN'], mode='lines+markers', name=name+pct_str, line=dict(dash='dot')))
                                        bar_labels.append(name)
                                        bar_values.append(pct)
                                        bar_colors.append('rgba(0, 200, 0, 0.7)' if pct >= 0 else 'rgba(255, 0, 0, 0.7)')
                                        
                                    if show_high: 
                                        pct, pct_str = calc_growth(df['HIGH'].iloc[0], df['HIGH'].iloc[-1])
                                        name = f"{label_base} (H)"
                                        fig_price.add_trace(go.Scatter(x=df['DATE'], y=df['HIGH'], mode='lines+markers', name=name+pct_str, line=dict(dash='dash')))
                                        bar_labels.append(name)
                                        bar_values.append(pct)
                                        bar_colors.append('rgba(0, 200, 0, 0.7)' if pct >= 0 else 'rgba(255, 0, 0, 0.7)')
                                        
                                    if show_low: 
                                        pct, pct_str = calc_growth(df['LOW'].iloc[0], df['LOW'].iloc[-1])
                                        name = f"{label_base} (L)"
                                        fig_price.add_trace(go.Scatter(x=df['DATE'], y=df['LOW'], mode='lines+markers', name=name+pct_str, line=dict(dash='dashdot')))
                                        bar_labels.append(name)
                                        bar_values.append(pct)
                                        bar_colors.append('rgba(0, 200, 0, 0.7)' if pct >= 0 else 'rgba(255, 0, 0, 0.7)')
                                        
                                    if show_close: 
                                        pct, pct_str = calc_growth(df['CLOSE'].iloc[0], df['CLOSE'].iloc[-1])
                                        name = f"{label_base} (C)"
                                        fig_price.add_trace(go.Scatter(x=df['DATE'], y=df['CLOSE'], mode='lines+markers', name=name+pct_str, line=dict(dash='solid')))
                                        bar_labels.append(name)
                                        bar_values.append(pct)
                                        bar_colors.append('rgba(0, 200, 0, 0.7)' if pct >= 0 else 'rgba(255, 0, 0, 0.7)')
                                
                                # --- 2. MULTI-DIMENSIONAL STRATEGY P&L GENERATION ---
                                entry_price = (0.125 * df['OPEN'].iloc[0] + 0.125 * df['CLOSE'].iloc[0] + 0.375 * df['HIGH'].iloc[0] + 0.375 * df['LOW'].iloc[0])
                                config['calculated_entry'] = entry_price
                                
                                if chart_type == "Candlestick Chart":
                                    if config["trade_direction"] == "Buy (Long)":
                                        pnl_open = (df['OPEN'] - entry_price) * config["lot_size"]
                                        pnl_high = (df['HIGH'] - entry_price) * config["lot_size"]
                                        pnl_low = (df['LOW'] - entry_price) * config["lot_size"]
                                        pnl_close = (df['CLOSE'] - entry_price) * config["lot_size"]
                                        df['P&L (Points)'] = df['CLOSE'] - entry_price
                                    else:
                                        pnl_open = (entry_price - df['OPEN']) * config["lot_size"]
                                        pnl_high = (entry_price - df['HIGH']) * config["lot_size"]
                                        pnl_low = (entry_price - df['LOW']) * config["lot_size"]
                                        pnl_close = (entry_price - df['CLOSE']) * config["lot_size"]
                                        df['P&L (Points)'] = entry_price - df['CLOSE']
                                        
                                    df['TOTAL P&L (₹)'] = df['P&L (Points)'] * config["lot_size"]
                                    df['PNL_OPEN'] = pnl_open
                                    df['PNL_HIGH'] = pnl_high
                                    df['PNL_LOW'] = pnl_low
                                    df['PNL_CLOSE'] = pnl_close
                                else:
                                    if config["trade_direction"] == "Buy (Long)":
                                        df['P&L (Points)'] = df['CLOSE'] - entry_price
                                        pnl_open = (df['OPEN'] - entry_price) * config["lot_size"]
                                        pnl_high = (df['HIGH'] - entry_price) * config["lot_size"]
                                        pnl_low = (df['LOW'] - entry_price) * config["lot_size"]
                                        pnl_close = (df['CLOSE'] - entry_price) * config["lot_size"]
                                    else:
                                        df['P&L (Points)'] = entry_price - df['CLOSE']
                                        pnl_open = (entry_price - df['OPEN']) * config["lot_size"]
                                        pnl_high = (entry_price - df['LOW']) * config["lot_size"]
                                        pnl_low = (entry_price - df['HIGH']) * config["lot_size"]
                                        pnl_close = (entry_price - df['CLOSE']) * config["lot_size"]
                                        
                                    df['TOTAL P&L (₹)'] = df['P&L (Points)'] * config["lot_size"]
                                    df['PNL_OPEN'] = pnl_open
                                    df['PNL_HIGH'] = pnl_high
                                    df['PNL_LOW'] = pnl_low
                                    df['PNL_CLOSE'] = pnl_close
                                
                                if chart_type == "Candlestick Chart":
                                    fig_pnl_individual.add_trace(go.Candlestick(
                                        x=df['DATE'],
                                        open=pnl_open,
                                        high=pnl_high,
                                        low=pnl_low,
                                        close=pnl_close,
                                        name=f"{label_base} ({action_short} P&L Candle)"
                                    ))
                                else:
                                    fig_pnl_individual.add_trace(go.Scatter(
                                        x=df['DATE'], 
                                        y=df['TOTAL P&L (₹)'], 
                                        mode='lines+markers', 
                                        name=f"{label_base} ({action_short} P&L)"
                                    ))
                                
                                df['SERIES_SOURCE'] = f"Block #{block_idx+1}"
                                df['STRIKE'] = strike
                                df['SYMBOL'] = config["symbol"]
                                df['OPTION_TYPE'] = config["option_type"]
                                df['TRADE_ACTION'] = action_short
                                df['ENTRY_PRICE_REF'] = f"Weighted ({entry_price:.2f})"
                                
                                all_combined_dfs.append(df[['DATE', 'SERIES_SOURCE', 'SYMBOL', 'STRIKE', 'OPTION_TYPE', 'TRADE_ACTION', 'OPEN', 'HIGH', 'LOW', 'CLOSE', 'PNL_OPEN', 'PNL_HIGH', 'PNL_LOW', 'PNL_CLOSE', 'TOTAL P&L (₹)']])
                                
                                data_found = True
                                break
                            else:
                                break
                                
                        except Exception as e:
                            if attempt < max_retries - 1:
                                time.sleep(1.5)
                            else:
                                st.warning(f"Failed to fetch data for {label_base}. Exact Error: {e}")
                    
                    time.sleep(1) 
                    
        # ==========================================
        # 4. OUTPUT RENDERING (CHARTS & EXPORT)
        # ==========================================
        if data_found:
            st.success("Successfully generated Analysis & Backtest!")
            
            master_df = pd.concat(all_combined_dfs, ignore_index=True)
            
            # --- CHART 1a: PREMIUM LINE/CANDLE CHART ---
            st.markdown("### 📈 1a. Multi-Series Premium Price Chart")
            
            fig_price.update_layout(
                title="<b>Historical Premium Movement</b>", 
                xaxis_title="Trading Date", 
                yaxis_title="Premium Price (₹)", 
                yaxis=dict(side="right"), 
                hovermode="x unified", 
                template="plotly_white",
                margin=dict(t=160), 
                xaxis_rangeslider_visible=False,
                legend=dict(
                    title=" ",
                    orientation="h",       
                    yanchor="bottom",      
                    y=1.10,             
                    xanchor="left",        
                    x=0,
                    bordercolor="lightgray",
                    borderwidth=1
                )
            )
            st.plotly_chart(fig_price, use_container_width=True)
            
            # --- CHART 1b: PREMIUM GROWTH BAR CHART ---
            st.markdown("### 📊 1b. Premium Growth Breakdown (Bar Chart)")
            
            fig_growth_bar.add_trace(go.Bar(
                x=bar_labels,
                y=bar_values,
                marker_color=bar_colors,
                text=[f"{val:+.1f}%" for val in bar_values],
                textposition='auto',
                width=0.3
            ))
            
            x_range_max = max(7.5, len(bar_labels) - 0.5)

            fig_growth_bar.update_layout(
                title="<b>Total Percentage Growth (Entry vs Exit)</b>",
                yaxis_title="Growth (%)",
                template="plotly_white",
                height=350,
                margin=dict(t=60, b=120, l=10, r=10),
                xaxis=dict(
                    tickangle=-45,
                    tickfont=dict(size=10),
                    range=[-0.5, x_range_max]
                )
            )
            st.plotly_chart(fig_growth_bar, use_container_width=True)

            # --- CHART 2: STRATEGY P&L CHART ---
            st.markdown("### 📊 2. Individual Strategy Profit & Loss (P&L)")
            entry_summary_text = " | ".join([f"**Block #{idx+1} Entry:** ₹{cfg.get('calculated_entry', 0):.2f}" for idx, cfg in enumerate(all_series_inputs)])
            st.markdown(f"<p style='color: gray; font-size: 14px; margin-top: -10px;'>📊 {entry_summary_text}</p>", unsafe_allow_html=True)

            fig_pnl_individual.add_hline(y=0, line_dash="dash", line_color="black", annotation_text="Breakeven (₹0)")
            
            fig_pnl_individual.update_layout(
                title="<b>Cumulative P&L (Broken Down by Strategy)</b>", 
                xaxis_title="Trading Date", 
                yaxis_title="Net P&L (₹)", 
                yaxis=dict(side="right"), 
                hovermode="x unified", 
                template="plotly_white", 
                margin=dict(t=160), 
                xaxis_rangeslider_visible=False,
                legend=dict(
                    title=" ",
                    orientation="h",       
                    yanchor="bottom",      
                    y=1.10,             
                    xanchor="left",        
                    x=0,
                    bordercolor="lightgray",
                    borderwidth=1
                )
            )
            st.plotly_chart(fig_pnl_individual, use_container_width=True)
            
            # ==========================================
            # 🏦 3. NET PORTFOLIO PROFIT & LOSS RANGE
            # ==========================================
            st.markdown("### 🏦 3. Net Portfolio Profit & Loss Range")
           
            portfolio_pnl = master_df.groupby('DATE').agg({
                'PNL_OPEN': 'sum',
                'PNL_HIGH': 'sum',
                'PNL_LOW': 'sum',
                'PNL_CLOSE': 'sum',
                'TOTAL P&L (₹)': 'sum'
            }).reset_index()

            portfolio_pnl['ROW_MAX'] = portfolio_pnl[['PNL_OPEN', 'PNL_HIGH', 'PNL_LOW', 'PNL_CLOSE']].max(axis=1)
            portfolio_pnl['ROW_MIN'] = portfolio_pnl[['PNL_OPEN', 'PNL_HIGH', 'PNL_LOW', 'PNL_CLOSE']].min(axis=1)

            overall_max = portfolio_pnl['ROW_MAX'].max()
            overall_min = portfolio_pnl['ROW_MIN'].min()
            final_closing_val = portfolio_pnl['TOTAL P&L (₹)'].iloc[-1] if not portfolio_pnl.empty else 0.0

            fig_portfolio = go.Figure()

            dates = portfolio_pnl['DATE'].tolist()
            maxs = portfolio_pnl['ROW_MAX'].tolist()
            mins = portfolio_pnl['ROW_MIN'].tolist()

            shade_dates = []
            shade_maxs = []
            shade_mins = []

            for i in range(len(dates)):
                shade_dates.append(dates[i])
                shade_maxs.append(maxs[i])
                shade_mins.append(mins[i])

                if i < len(dates) - 1:
                    if (maxs[i] > 0 and maxs[i+1] < 0) or (maxs[i] < 0 and maxs[i+1] > 0):
                        ratio = abs(maxs[i]) / (abs(maxs[i]) + abs(maxs[i+1]))
                        z_date = dates[i] + ((dates[i+1] - dates[i]) * ratio)
                        z_max = 0
                        z_min = mins[i] + ((mins[i+1] - mins[i]) * ratio)
                        shade_dates.append(z_date)
                        shade_maxs.append(z_max)
                        shade_mins.append(z_min)

                    if (mins[i] > 0 and mins[i+1] < 0) or (mins[i] < 0 and mins[i+1] > 0):
                        ratio = abs(mins[i]) / (abs(mins[i]) + abs(mins[i+1]))
                        z_date = dates[i] + ((dates[i+1] - dates[i]) * ratio)
                        z_min = 0
                        z_max = maxs[i] + ((maxs[i+1] - maxs[i]) * ratio)
                        shade_dates.append(z_date)
                        shade_maxs.append(z_max)
                        shade_mins.append(z_min)

            df_shade = pd.DataFrame({'DATE': shade_dates, 'MAX': shade_maxs, 'MIN': shade_mins})
            df_shade = df_shade.sort_values('DATE').drop_duplicates(subset=['DATE']).reset_index(drop=True)
            safe_dates = df_shade['DATE']

            pos_upper = df_shade['MAX'].apply(lambda x: x if x > 0 else 0)
            pos_lower = df_shade['MIN'].apply(lambda x: x if x > 0 else 0)
            neg_upper = df_shade['MAX'].apply(lambda x: x if x < 0 else 0)
            neg_lower = df_shade['MIN'].apply(lambda x: x if x < 0 else 0)

            fig_portfolio.add_trace(go.Scatter(x=safe_dates, y=pos_lower, mode='lines', line=dict(width=0), showlegend=False, hoverinfo='skip'))
            fig_portfolio.add_trace(go.Scatter(x=safe_dates, y=pos_upper, mode='none', fill='tonexty', fillcolor='rgba(0, 200, 0, 0.25)', showlegend=False, hoverinfo='skip'))

            fig_portfolio.add_trace(go.Scatter(x=safe_dates, y=neg_upper, mode='lines', line=dict(width=0), showlegend=False, hoverinfo='skip'))
            fig_portfolio.add_trace(go.Scatter(x=safe_dates, y=neg_lower, mode='none', fill='tonexty', fillcolor='rgba(255, 0, 0, 0.25)', showlegend=False, hoverinfo='skip'))

            hover_texts_max = [f"<b>Date:</b> {d.strftime('%Y-%m-%d')}<br><b>Max High:</b> ₹{h:.2f}" for d, h in zip(portfolio_pnl['DATE'], portfolio_pnl['ROW_MAX'])]
            hover_texts_min = [f"<b>Date:</b> {d.strftime('%Y-%m-%d')}<br><b>Min Low:</b> ₹{l:.2f}" for d, l in zip(portfolio_pnl['DATE'], portfolio_pnl['ROW_MIN'])]

            fig_portfolio.add_trace(go.Scatter(
                x=portfolio_pnl['DATE'],
                y=portfolio_pnl['ROW_MAX'],
                mode='lines+markers',
                name='Absolute Max High Boundary',
                line=dict(color='green', width=2),
                text=hover_texts_max, hovertemplate="%{text}<extra></extra>"
            ))

            fig_portfolio.add_trace(go.Scatter(
                x=portfolio_pnl['DATE'],
                y=portfolio_pnl['ROW_MIN'],
                mode='lines+markers',
                name='Absolute Min Low Boundary',
                line=dict(color='red', width=2),
                text=hover_texts_min, hovertemplate="%{text}<extra></extra>"
            ))

            if chart_type != "Candlestick Chart":
                pnl_colors = ['green' if val >= 0 else 'red' for val in portfolio_pnl['TOTAL P&L (₹)']]
                fig_portfolio.add_trace(go.Scatter(
                    x=portfolio_pnl['DATE'],
                    y=portfolio_pnl['TOTAL P&L (₹)'],
                    mode='lines+markers',
                    name="NET PORTFOLIO P&L",
                    line=dict(color='gray', width=2, dash='dot'),
                    marker=dict(color=pnl_colors, size=8, line=dict(width=1, color='white'))
                ))

            fig_portfolio.add_hline(y=0, line_dash="solid", line_width=2, line_color="black", annotation_text="Breakeven (₹0)")
            fig_portfolio.update_layout(
                title=f"<b>Combined Portfolio Range (Max High = {overall_max:.1f} , Min Low = {overall_min:.1f}, Final Closing P&L = {final_closing_val:.1f})</b>",
                xaxis_title="Trading Date",
                yaxis_title="P&L Extreme Range (₹)",
                yaxis=dict(side="right"),
                hovermode="x unified" if chart_type != "Candlestick Chart" else "closest",
                template="plotly_white",
                xaxis_rangeslider_visible=False,
                showlegend=False
            )
            st.plotly_chart(fig_portfolio, use_container_width=True)
            
            # --- DISPLAY 3 METRIC BOXES BELOW THE GRAPH ---
            st.markdown("#### 📌 Key Portfolio Performance Summary")
            m_col1, m_col2, m_col3 = st.columns(3)
            with m_col1:
                st.metric(label="📈 Absolute Max High", value=f"₹{overall_max:,.2f}")
            with m_col2:
                st.metric(label="📉 Absolute Min Low", value=f"₹{overall_min:,.2f}")
            with m_col3:
                st.metric(label="🏁 Final Closing P&L", value=f"₹{final_closing_val:,.2f}", delta=f"{final_closing_val:+,.2f}")
            
            with st.expander("View Consolidated Daily Ledger & Trade Data"):
                st.dataframe(master_df.sort_values(by=['DATE', 'SERIES_SOURCE', 'STRIKE'], ascending=[False, True, True]))

            # ==========================================
            # 5. EXPORT HUB
            # ==========================================
            st.markdown("---")
            st.markdown("### 💾 Export Analysis & Reports")
            
            dl_col1, dl_col2 = st.columns(2)

            with dl_col1:
                csv_data = master_df.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="📊 Download Raw Data Ledger (CSV)",
                    data=csv_data,
                    file_name=f"options_portfolio_ledger_{today}.csv",
                    mime="text/csv",
                    use_container_width=True
                )

            with dl_col2:
                try:
                    with st.spinner("Generating Microsoft Word Document..."):
                        doc = Document()
                        doc.add_heading(f'Options Strategy Executive Report - {today}', 0)
                        doc.add_paragraph("This report contains the automated backtest results and historical premium charts for the configured options portfolio.")

                        img1_bytes = fig_price.to_image(format="png", width=900, height=500, scale=1.5)
                        img1b_bytes = fig_growth_bar.to_image(format="png", width=900, height=350, scale=1.5)
                        img2_bytes = fig_pnl_individual.to_image(format="png", width=900, height=500, scale=1.5)
                        img3_bytes = fig_portfolio.to_image(format="png", width=900, height=500, scale=1.5)

                        doc.add_heading('1a. Multi-Series Premium Price Chart', level=1)
                        doc.add_picture(io.BytesIO(img1_bytes), width=Inches(6.5))
                        
                        doc.add_heading('1b. Premium Growth Breakdown', level=1)
                        doc.add_picture(io.BytesIO(img1b_bytes), width=Inches(6.5))
                        
                        doc.add_page_break()

                        doc.add_heading('2. Individual Strategy Profit & Loss (P&L)', level=1)
                        doc.add_picture(io.BytesIO(img2_bytes), width=Inches(6.5))

                        doc.add_heading('3. Net Portfolio Profit & Loss Range', level=1)
                        doc.add_picture(io.BytesIO(img3_bytes), width=Inches(6.5))

                        doc_buffer = io.BytesIO()
                        doc.save(doc_buffer)
                        doc_bytes = doc_buffer.getvalue()
                        
                        st.download_button(
                            label="📝 Download as Microsoft Word (.docx)",
                            data=doc_bytes,
                            file_name=f"Options_Article_Draft_{today}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                except ValueError as e:
                    st.error("⚠️ To download Word docs, run `py -m pip install kaleido` in your terminal, and restart.")
                
        else:
            st.error("Could not find data matches for any of the configured options. Check your dates.")

# ==========================================
# CUSTOM HTML (SCROLL TO TOP)
# ==========================================
st.markdown(
    """
    <style>
        .scroll-to-top {
            position: fixed;
            bottom: 30px;
            left: 30px;
            width: 50px;
            height: 50px;
            background-color: #2e3b4e;
            color: white !important;
            border-radius: 50%;
            display: flex;
            justify-content: center;
            align-items: center;
            font-size: 24px;
            text-decoration: none;
            box-shadow: 2px 2px 6px rgba(0,0,0,0.4);
            z-index: 99999;
            transition: background-color 0.3s, transform 0.3s;
        }
        .scroll-to-top:hover {
            background-color: #1a2333;
            transform: scale(1.1);
        }
    </style>
    <a href="#top" class="scroll-to-top" title="Scroll to Top">⬆️</a>
    """,
    unsafe_allow_html=True
)
